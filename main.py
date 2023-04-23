import os
import time
from datetime import date
import re
import requests
from bs4 import BeautifulSoup
import spacy
import torch
import numpy as np
from scipy.spatial.distance import cosine
from transformers import AutoTokenizer, AutoModel, T5Tokenizer, T5ForConditionalGeneration, pipeline
from newspaper import Article
from sentence_transformers import SentenceTransformer
from googleapiclient.discovery import build
import PyPDF4
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
import pdfplumber
from citeproc import CitationStylesStyle, CitationStylesBibliography, Citation, CitationItem, formatter
from citeproc.source.json import CiteProcJSON
import random
import json
from sklearn.cluster import DBSCAN
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.preprocessing import StandardScaler
from typing import Tuple, List, Dict, Union
from collections import defaultdict
import matplotlib.pyplot as plt
plt.ion()
from sklearn.manifold import TSNE
import umap
import matplotlib.pyplot as plt
import seaborn as sns

from threading import Lock

# Global dictionary to store progress information for each request
progress = {}
progress_lock = Lock()

def update_progress(request_id, stage, num, outof):
    with progress_lock:
        progress[request_id] = {'stage': stage, 'progress': f"{num} / {outof}", 'as_num': num/outof, 'num': num, 'outof': outof}


startTime = time.time()
times  = []
global weightedTimeTotal
weightedTimeTotal = 0
global urlTimeTotal
global totalUrls
urlTimeTotal = 0
totalUrls = 0

class ContextSentence:
    def __init__(self, text: str, similarity: float):
        self.text = text
        self.similarity = similarity

class RelevantSentence:
    def __init__(self, sentence: str, is_relevant: bool, before_context: List[ContextSentence], after_context: List[ContextSentence]):
        self.text = sentence
        self.is_relevant = is_relevant
        self.before_context = before_context
        self.after_context = after_context

class Evidence:
    def __init__(self, url: str, tagline: str, relevant_sentences: List[RelevantSentence]):
        self.url = url
        self.tagline = tagline
        self.relevant_sentences = relevant_sentences

api_key = os.environ.get('API_KEY')
CSE = os.environ.get('CSE')

nlp = spacy.load('en_core_web_lg')
sbert_model = SentenceTransformer('paraphrase-distilroberta-base-v2')
ner_pipeline = pipeline("ner", model="dbmdz/bert-large-cased-finetuned-conll03-english")

tokenizer = T5Tokenizer.from_pretrained("t5-small")
model = T5ForConditionalGeneration.from_pretrained("t5-base")

def generate_tagline(text: str) -> str:
    inputs = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=1024, truncation=True)
    outputs = model.generate(inputs, max_length=50, min_length=10, length_penalty=2.0, num_beams=4, early_stopping=True, top_p=0.9)
    return tokenizer.decode(outputs[0], skip_special_tokens=True)

def get_sentence_embedding_sbert(sentence: str):
    sentence_embedding = sbert_model.encode(sentence, convert_to_tensor=True)
    return sentence_embedding

def sbert_cosine_similarity(sentence1: str, sentence2: str) -> float:
    embedding1 = get_sentence_embedding_sbert(sentence1)
    embedding2 = get_sentence_embedding_sbert(sentence2)
    similarity = 1 - cosine(embedding1, embedding2)
    return similarity

def extract_article_content(url: str) -> str:
    article = Article(url)
    article.download()
    article.parse()

    cleaned_authors = []

    if article.authors:
        cleaned_authors = [extract_author_name(author) for author in article.authors]
        cleaned_authors = [author for author in cleaned_authors if author != ""]
        author_part = " & ".join(cleaned_authors)
    else:
        author_part = "Unknown Author"

    year_part = ("'"+article.publish_date.strftime("%Y")) if article.publish_date else "Unknown Year"
    citation = f"{author_part} {year_part}"

    if flag_for_manual_review(article.text, cleaned_authors):
        citation += " (manual review recommended)"

    return article.text, citation

def extract_author_name(author_info: str) -> str:
    ner_results = ner_pipeline(author_info)
    names = [result["word"] for result in ner_results if result["entity"] == "B-PER" or result["entity"] == "I-PER"]
    name = " ".join(names).replace(" ##", "")
    return name if name else ""

def flag_for_manual_review(text: str, authors: List[str]) -> bool:
    # Set a threshold for the number of authors that warrant manual review
    multiple_authors_threshold = 3

    if len(authors) >= multiple_authors_threshold:
        return True

    # Check if there are any other patterns indicating multiple authors
    multiple_authors_patterns = [
        r'\band\b',  # Look for the word "and" between names
        r',',  # Look for commas between names
    ]

    for pattern in multiple_authors_patterns:
        if re.search(pattern, text):
            return True

    return False

def preprocess_text(text: str) -> str:
    # Create a spaCy doc object
    doc = nlp(text)
    
    # Preprocess tokens: lowercase, remove stopwords and punctuation, and lemmatize
    preprocessed_tokens = [
        token.lemma_.lower()
        for token in doc
        if not token.is_stop and not token.is_punct
    ]
    
    # Join the preprocessed tokens to create the preprocessed text
    preprocessed_text = ' '.join(preprocessed_tokens)
    
    return preprocessed_text

def get_article_content(url: str) -> str:
    _, file_extension = os.path.splitext(url)
    file_extension = file_extension.lower()

    user_agent_list = [
        'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3',
        'Mozilla/5.0 (Windows NT 6.1; WOW64; rv:54.0) Gecko/20100101 Firefox/54.0',
        'Mozilla/5.0 (Windows NT 6.1; WOW64; Trident/7.0; AS; rv:11.0) like Gecko',
        'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/64.0.3282.140 Safari/537.36 Edge/17.17134'
    ]
    user_agent = random.choice(user_agent_list)

    headers = {'User-Agent': user_agent}

    # Download the file
    response = requests.get(url, headers=headers, timeout=30)
    response.raise_for_status()
    
    cite = "Citation Unavailable"

    # If the file is a PDF
    if file_extension == '.pdf':
        content = extract_pdf_content_from_response(response)
        content = preprocess_pdf_text(content)

    # If the file is a Word document
    elif file_extension in ['.doc', '.docx']:
        content = extract_word_document_content_from_response(response)

    # If the file is a text file
    elif file_extension == '.txt':
        content = extract_text_file_content_from_response(response)

    # If the file is an RTF file
    elif file_extension == '.rtf':
        content = extract_word_document_content_from_response(response)

    # If the file is HTML or an unsupported type
    else:
        content, cite = extract_article_content(url)

    return content, cite


def extract_pdf_content_from_response(response: requests.Response) -> str:    
    with io.BytesIO(response.content) as f:
        return extract_pdf_content(f)

def extract_word_document_content_from_response(response: requests.Response) -> str:
    with io.BytesIO(response.content) as f:
        return extract_word_document_content(f)

def extract_text_file_content_from_response(response: requests.Response) -> str:
    return response.text
    
def extract_word_document_content(file_obj: io.BytesIO) -> str:
    doc = Document(file_obj)
    content = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    return content


def find_relevant_sentences(text:str, query:str, context:int=4, similarity_threshold:float=0.5) -> Tuple[Tuple[str, bool, List[Tuple[str,float]], List[Tuple[str,float]], int, int], str]:
    start_time = time.time()
    global weightedTimeTotal

    if(len(text)>1000000):
        raise Exception("Text is too long")

    doc = nlp(text)
    query_doc = nlp(query)
    relevant_sentences = []
    sentences = list(doc.sents)
    included_in_context = set()
    char_index = 0

    for i, sentence in enumerate(sentences):

        if not is_informative(sentence):
            char_index += len(sentence.text) + 1
            continue

        text = preprocess_text(sentence.text)
        if(len(text)==0):
            char_index += len(sentence.text) + 1
            continue

        similarity = sbert_cosine_similarity(text, query)



        if re.search(r'(\d+)?(\.\d+)?( ?million| ?billion| ?trillion| ?percent| ?%)', sentence.text, flags=re.IGNORECASE):
            similarity *= 2
        elif re.search(r'\d|%', sentence.text):
            similarity *= 1.1
        elif re.search(r'because|since|so', sentence.text):
            similarity *= 1.25
        if contains_named_entities(sentence):
            similarity *= 1.25

        if similarity > similarity_threshold:
            start_index_char = char_index
            end_index_char = char_index + len(sentence.text)
            if i not in included_in_context:
                start_index = max(0, i - context)
                end_index = min(len(sentences), i + context + 1)
                prev_context = [(sentences[j].text.strip(), sbert_cosine_similarity(sentences[j].text, sentence.text)) for j in range(start_index, i)]
                next_context = [(sentences[j].text.strip(), sbert_cosine_similarity(sentences[j].text, sentence.text)) for j in range(i + 1, end_index)]

                # Add the sentences from prev_context and next_context to the set
                for j in range(start_index, end_index):
                    if j != i:
                        included_in_context.add(j)

                relevant_sentences.append((sentence, True, prev_context, next_context, start_index_char, end_index_char))
            else:
                # If the sentence is already part of the context, mark it as relevant without changing the context
                for j, (rel_sentence, is_relevant, prev, after, _, _) in enumerate(relevant_sentences):
                    if rel_sentence == sentence:
                        relevant_sentences[j] = (rel_sentence, True, prev, after, start_index_char, end_index_char)
                        break

        char_index += len(sentence.text) + 1

    relevant_text = [sentence.text for sentence, _, _, _, _, _ in relevant_sentences]

    delta_t = time.time() - start_time
    weightedTimeTotal += delta_t / len(sentences)

    return relevant_sentences, relevant_text

def is_informative(sentence):
    text = sentence.text.strip()
    token_count = len(sentence)

    if token_count <= 2:
        return False

    if re.match(r'^\s*(chapter|section|introduction|conclusion|acknowledgment|reference|table of contents)\s*$', text, flags=re.IGNORECASE):
        return False

    return True
def contains_named_entities(sentence):
    return any([token.ent_type_ for token in sentence])


def generate_queries(topic:str, side:str, argument:str) -> List[str]:
    side_keywords = {
        'pro': ['advantages', 'benefits', 'positive aspects', 'strengths'],
        'con': ['disadvantages', 'drawbacks', 'negative aspects', 'weaknesses']
    }

    # Combine topic, side, and argument with relevant keywords
    queries = []
    for keyword in side_keywords[side]:
        queries.append(f"{topic} {keyword} {argument}")
        queries.append(f"{topic} {argument} {keyword}")

    # Add queries that focus only on the argument
    queries.append(f"{topic} {argument}")
    queries.append(f"{argument} {topic}")

    return queries

#Generate query from topic side and argument provided
def generate_query(topic: str, side: str, argument: str) -> str:
    side_keywords = {
        'pro': ['advantages', 'benefits', 'positive aspects', 'strengths'],
        'con': ['disadvantages', 'drawbacks', 'negative aspects', 'weaknesses'],
        'sup': ['']
    }
    
    # Combine topic, side, and argument with relevant keywords
    side_kw_string = ' '.join(side_keywords[side])
    
    # Assign different weights to topic, side, and argument
    query = f"{topic} {topic} {side_kw_string} {side_kw_string} {argument} {argument}"
    
    return query

def generate_search_query(topic:str, side:str, argument:str) -> str: 
    query = f"{argument}"
    return query

#extract text from pdf
def extract_pdf_content(file_obj: io.BytesIO) -> str:
    def group_chars_by_x(characters, threshold=1.0):
        groups = []
        current_group = []
        prev_x = None
        for char in characters:
            if prev_x is not None and char["x0"] - prev_x > threshold:
                groups.append(current_group)
                current_group = []
            current_group.append(char)
            prev_x = char["x1"]
        groups.append(current_group)
        return groups

    with pdfplumber.open(file_obj) as pdf:
        text = []

        for page in pdf.pages:
            current_line = []
            current_y = -1

            for char in sorted(page.chars, key=lambda x: (x["top"], x["x0"])):
                if char["top"] != current_y:
                    if current_line:
                        words = group_chars_by_x(current_line)
                        text.append(" ".join("".join(c["text"] for c in word) for word in words))
                        current_line = []
                    current_y = char["top"]

                
                current_line.append(char)

            if current_line:
                words = group_chars_by_x(current_line)
                text.append(" ".join("".join(c["text"] for c in word) for word in words))

        return "\n".join(text)

def search_articles(query: str, api_key: str, CSE: str, num_results: int=10):
    service = build("customsearch", "v1", developerKey=api_key)
    urls = []
    start_index = 1

    while len(urls) < num_results:
        response = service.cse().list(q=query, cx=CSE, start=start_index, fileType="pdf").execute()
        results = response.get('items', [])

        for result in results:
            url = result.get('link')
            if url:
                urls.append(url)

        # Handle pagination
        start_index += 10
        if start_index > 100 or not response.get('queries').get('nextPage'):
            break

    return urls[:num_results]

import re

def preprocess_pdf_text(text: str):
    # Remove non-alphanumeric characters, except for spaces, periods, and commas
    cleaned_text = re.sub(r"[^a-zA-Z0-9,. ]", " ", text)

    # Remove multiple spaces
    cleaned_text = re.sub(r"\s+", " ", cleaned_text)

    # Remove multiple periods
    cleaned_text = re.sub(r"\.+", ".", cleaned_text)

    # Remove multiple commas
    cleaned_text = re.sub(r",+", ",", cleaned_text)

    # Remove spaces before commas and periods
    cleaned_text = re.sub(r"\s+([,.])", r"\1", cleaned_text)

    cleaned_text = re.sub(r'\n+', '\n', cleaned_text)

    # Remove extra spaces at the beginning and end of the text
    cleaned_text = cleaned_text.strip()
    
    return cleaned_text

def main(topic: str, side: str, argument: str, num_results: int = 10, request_id=None) -> Dict[str, List[Dict[str, List[RelevantSentence]]]]:
    url_sentence_map = defaultdict(list)
    raw_data = {}
    relevant_sentences:List[List[Tuple[str, bool, List[Tuple(str, float)], List[Tuple(str, float)]]]] = []
    resulting_sentences:List[Evidence] = []
    global totalUrls
    global urlTimeTotal
    if side == "sup":
        topic = " "
    query = generate_query(topic, side, argument)
    processed_query = preprocess_text(query)
    sQuery = generate_search_query(topic, side, argument)
    processed_sQuery = preprocess_text(sQuery)

    urls = search_articles(processed_sQuery, api_key, CSE, num_results=num_results)
    url_text_map = {}

    print(urls)
    total_urls = len(urls)

    #Threading for backend
    if request_id is not None:
        update_progress(request_id, "fetching", 0, total_urls)  # Initialize progress to 0

    curURL = 0
    for url in urls:
        try:
            #logging
            urlTimeStart = time.time()

            #extract url content
            content, cite = get_article_content(url)

            # Update progress percentage
            update_progress(request_id, "fetching", (curURL + 1), total_urls)

            #logging
            urlDeltaTime = time.time() - urlTimeStart
            urlTimeTotal += urlDeltaTime
            totalUrls += 1

            if content:
                url_text_map[url] = content, cite
        except Exception as e:
            print(f"Error fetching URL {url}: {e}")

    curURL = 0

    for url, (text, citation) in url_text_map.items():

        individualStart = time.time()
        try:
            if request_id is not None:

                # Update progress percentage
                update_progress(request_id, "processing", (curURL + 1), total_urls)

            res_sentence, relevant_text = find_relevant_sentences(text, query)
        except Exception as e:
            print(f"Error finding relevant content in {url}: {e}")
            continue
        deltaTime = time.time() - individualStart
        times.append(deltaTime)

        rel_sentences = {}

        for i, (sentence, is_relevant, prev, after, start, end) in enumerate(res_sentence):
            rel_sentences[i] = (sentence, is_relevant, prev, after, start, end)

        # Perform clustering on the relevant sentences
        clusters_indices, representative_sentences = cluster_relevant_sentences([sent.text for sent, _, _, _, _, _ in res_sentence])

        # Generate taglines for each cluster
        taglines = [generate_tagline(rep_sentence) for rep_sentence in representative_sentences]

        print(rel_sentences.keys())
        # Store the information in the url_sentence_map, using the new structure
        for i in range(len(clusters_indices)):
            # Convert Span objects to strings for each sentence in the cluster
            true_sentence = [(rel_sentences[j][0].text, rel_sentences[j][1], rel_sentences[j][2], rel_sentences[j][3]) for j in clusters_indices[i]]

            url_sentence_map[url].append({
                'tagline': taglines[i],
                'relevant_sentences': true_sentence,
                'citation': citation
            })
        raw_data[url] = {'full_text': relevant_text, 'prompt': query}
        print(f"Finished processing URL: {url}, Relevant Sentences: {len(res_sentence)}, URL Number: {curURL}")
        curURL += 1


    del progress[request_id]  # Remove the request progress on completion
    return url_sentence_map, raw_data


def cluster_relevant_sentences(sentences: List[str], eps: float = 0.22, min_samples: int = 2, outlier_eps: float = 0.7):
    if not sentences:
        return [], []

    # Generate sentence embeddings using Universal Sentence Encoder
    embeddings = sbert_model.encode(sentences)

    # Perform DBSCAN clustering
    dbscan = DBSCAN(eps=eps, min_samples=min_samples, metric='cosine')
    dbscan.fit(embeddings)

    # Group sentence indices by cluster
    unique_labels = set(dbscan.labels_)
    num_clusters = len(unique_labels) - (1 if -1 in unique_labels else 0)
    clustered_indices = [[] for _ in range(num_clusters)]
    for idx, label in enumerate(dbscan.labels_):
        if label != -1:
            clustered_indices[label].append(idx)

    # Combine all sentences in each cluster
    representative_sentences = []
    for label in unique_labels:
        if label != -1:
            cluster_indices = np.where(dbscan.labels_ == label)[0]
            cluster_sentences = [sentences[idx] for idx in cluster_indices]
            combined_sentence = ' '.join(cluster_sentences)
            representative_sentences.append(combined_sentence)

    # Perform a second pass of DBSCAN on the outliers
    outlier_sentences = [sentences[i] for i, label in enumerate(dbscan.labels_) if label == -1]
    if outlier_sentences:
        outlier_embeddings = sbert_model.encode(outlier_sentences)
        outlier_dbscan = DBSCAN(eps=outlier_eps, min_samples=1, metric='cosine')
        outlier_dbscan.fit(outlier_embeddings)

        # Include clustered outlier sentences in the `clustered_indices` as seperate clusters
        for idx, label in enumerate(outlier_dbscan.labels_):
            original_idx = sentences.index(outlier_sentences[idx])
            if label != -1:
                new_label = num_clusters + label
                if new_label < len(clustered_indices):
                    clustered_indices[new_label].append(original_idx)
                else:
                    clustered_indices.append([original_idx])
                    representative_sentences.append(outlier_sentences[idx])
            else:
                new_cluster_index = len(clustered_indices)
                clustered_indices.append([original_idx])
                representative_sentences.append(outlier_sentences[idx])

    return clustered_indices, representative_sentences



def apply_style(paragraph, font_size, bold=False, underline=False):
    run = paragraph.runs[0]
    run.font.size = Pt(font_size)
    run.bold = bold
    run.underline = underline

def add_table_of_contents(doc, current_url_map):

    # Add "Table of Contents" title
    para = doc.add_paragraph("Table of Contents", style='Title')
    apply_style(para, font_size=16, bold=True)
    doc.add_paragraph("\n")

    # Add tags of each cluster
    for url, clusters in current_url_map.items():
        for cluster in clusters:
            tagline = cluster['tagline']
            relevant_sentences = cluster['relevant_sentences']
            if len(relevant_sentences) > 0:
                para = doc.add_paragraph(tagline, style='Heading 1')
    
    doc.add_page_break()


def get_mla_citation(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        title = soup.title.string if soup.title else url

        reference = {
            'id': '1',
            'type': 'webpage',
            'title': title,
            'URL': url,
            'accessed': {'date-parts': [[date.fromtimestamp(time.time()).year, date.fromtimestamp(time.time()).month, date.fromtimestamp(time.time()).day]]}
        }

        bib_source = CiteProcJSON([reference])
        bib_style = CitationStylesStyle('modern-language-association', locale='en-US')
        bibliography = CitationStylesBibliography(bib_style, bib_source, formatter.plain)

        citation = Citation([CitationItem('1')])
        bibliography.register(citation)

        return bibliography.bibliography()[0]
    except Exception as e:
        print(f"Error generating MLA citation for URL {url}: {e}")
        return url

def write_sentences_to_word_doc(file_path: str, url_sentence_map: Dict[str, List[Dict[str, List[RelevantSentence]]]], info: Tuple[str, str]):

    num_urls_per_doc = 100
    docNum = 0
    url_count = 0

    while url_count < len(url_sentence_map):
        current_url_map = dict(list(url_sentence_map.items())[url_count:url_count + num_urls_per_doc])
        doc = Document()
        (topic, side, argument) = info
        doc.add_paragraph("Topic: " + topic + "\nSide: " + side + "\n" + "Arg: " + argument, style='Title')
        doc.add_page_break()

        add_table_of_contents(doc, current_url_map)

        for url, clusters in current_url_map.items():
            print(url+ " has " + str(len(clusters)) + " clusters")
            for cluster in clusters:
                tagline = cluster['tagline']
                relevant_sentences = cluster['relevant_sentences']

                if len(relevant_sentences) == 0:
                    url_count += 1
                    continue

                para = doc.add_paragraph("Tagline: " + tagline, style='Heading 1')
                apply_style_run(para.runs[0], font_size=14, bold=True, underline=True)
                doc.add_paragraph(url, style='Heading 2')

                url_para = doc.add_paragraph()
                for sentence, is_relevant, before_context, after_context in relevant_sentences:
                    if is_relevant:
                        r = url_para.add_run(sentence.strip() + " ")
                        apply_style_run(r, font_size=12, bold=True)

                    for context_sentence, context_similarity in before_context:
                        r = url_para.add_run(context_sentence.strip() + " ")
                        if context_similarity > 0.5:
                            apply_style_run(r, font_size=12, underline=True)
                        else:
                            apply_style_run(r, font_size=7)

                    for context_sentence, context_similarity in after_context:
                        r = url_para.add_run(context_sentence.strip() + " ")
                        if context_similarity > 0.5:
                            apply_style_run(r, font_size=12, underline=True)
                        else:
                            apply_style_run(r, font_size=7)

                url_count += 1
                print("Finished writing URL: " + url)

        doc.save(file_path + str(docNum) + ".docx")
        docNum += 1

def apply_style_run(run, font_size=None, bold=False, underline=False):
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.bold = bold
    if underline:
        run.underline = WD_UNDERLINE.SINGLE

import json

def save_to_json(file_path:str, url_sentence_map:dict, info:Tuple[str, str, str]):
    data = []
    (topic, side, argument) = info
    for url, clusters in url_sentence_map.items():
        for cluster in clusters:
            tagline = cluster['tagline']
            relevant_sentences = cluster['relevant_sentences']
            
            if len(relevant_sentences) == 0:
                continue

            temp_data = {"data": [], "tagline": tagline, 'topic': topic, 'side': side, 'argument': argument}

            for sentence, is_relevant, prev_context, next_context in relevant_sentences:
                sentence_data = {
                    'relevant_sentence': sentence,
                    'is_relevant': is_relevant,
                    'prev_context': [{"text": t, "similarity": s} for t, s in prev_context],
                    'next_context': [{"text": t, "similarity": s} for t, s in next_context],
                }
                temp_data['data'].append(sentence_data)
            
            temp_data['url'] = url
            data.append(temp_data)

    with open(file_path, 'w', encoding='utf-8') as json_file:
        json.dump(data, json_file, ensure_ascii=False, indent=4)

def visualize_embeddings(embeddings, labels):
    reducer = umap.UMAP()
    embeddings_2d = reducer.fit_transform(embeddings)

    plt.figure(figsize=(12, 8))
    sns.scatterplot(
        x=embeddings_2d[:, 0], y=embeddings_2d[:, 1],
        hue=labels,
        palette=sns.color_palette("hls", len(set(labels))),
        legend="full",
        alpha=0.7,
    )
    plt.title('UMAP visualization of sentence embeddings')
    plt.show()


topic = "Should not ban the collection of personal data through biometric recognition technology"
'''
arguments = [
    {
        "side": "pro",
        "argument": "Biometric technology is necessary for national security"
    },
    {
        "side": "pro",
        "argument": "Biometric technology is necessary for US hegemony"
    },
    {
        "side": "pro",
        "argument": "Biometric technology is necessary for banks"
    },
    {
        "side": "con",
        "argument": "Biometric techonolgy is not accurate"
    },
    {
        "side": "con",
        "argument": "Biometric technology allows identity theft"
    },
    {
        "side": "con",
        "argument": "Biometric tecohoogy used by authoritarian governments"
    },
    {
        "side": "con",
        "argument": "Biometric technology is necessary for the metaverse"
    },
]
'''
if __name__ == "__main__":

    arguments = [
        {
            "side": "sup",
            "argument": "nuclear war kills millions"
        }
    ]    

    for item in arguments:
        side = item["side"]
        argument = item["argument"]
        file_path = side+"_"+("_".join(argument.split(" ")[-3:]))
        url_sentence_map, raw_data = main(topic, side, argument, 4)

        write_sentences_to_word_doc(file_path, url_sentence_map, (topic, side, argument))
        save_to_json(file_path+".json", url_sentence_map, (topic, side, argument))

    timeElapsed = time.time() - startTime
    print("\nTIME: "+str(timeElapsed))
    print("AVERAGE TIME: "+str(sum(times)/len(times)))
    print("Weighted time: " + str(weightedTimeTotal/len(times)))
    print("Average query time: " + str(urlTimeTotal/totalUrls))

